#!/usr/bin/env python3
"""
Wattpad Episode Composer for The Boy Who Knew Me First Series

Generates DOCX files with formatting and images for easy copy-paste into Wattpad.
DOCX preserves bold, italic, and inline images when pasted.

Usage:
    uv run scripts/compose_for_wattpad.py <episode-path>

Examples:
    # Generate separate chapter DOCX files for Wattpad
    uv run scripts/compose_for_wattpad.py series/hockey-shuttle/season-01/episode-01-returning-to-center-ice
"""

import argparse
import sys
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def parse_args():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Generate DOCX files for Wattpad publishing with images and formatting",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument(
        "episode_path",
        type=str,
        help="Path to episode directory (e.g., series/hockey-shuttle/season-01/episode-01-returning-to-center-ice)"
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default="output/wattpad",
        help="Output directory relative to repo root. Default: output/wattpad"
    )
    return parser.parse_args()


def find_repo_root():
    """Find the repository root directory."""
    current = Path.cwd()
    while current != current.parent:
        if (current / ".git").exists() or (current / "series").exists():
            return current
        current = current.parent
    return Path.cwd()


def extract_episode_info(episode_path: Path):
    """Extract episode information from path."""
    episode_name = episode_path.name
    match = re.match(r"episode-(\d+)-(.+)", episode_name)

    if match:
        episode_num = int(match.group(1))
        episode_title = match.group(2).replace("-", " ").title()
    else:
        episode_num = "1"
        episode_title = episode_name

    return {
        "series": "The Boy Who Knew Me First",
        "season": "1",
        "episode": episode_num,
        "title": episode_title,
        "full_title": f"The Boy Who Knew Me First - Season 1, Episode {episode_num}: {episode_title}",
        "episode_slug": episode_name
    }


def read_chapter_files(episode_path: Path):
    """Read all chapter markdown files."""
    draft_dir = episode_path / "draft"

    if not draft_dir.exists():
        print(f"❌ Error: Draft directory not found at {draft_dir}")
        return []

    chapter_files = sorted(draft_dir.glob("chapter-*.md"))

    if not chapter_files:
        print(f"❌ Error: No chapter files found in {draft_dir}")
        return []

    chapters = []
    for chapter_file in chapter_files:
        try:
            content = chapter_file.read_text()

            # Extract chapter title
            title = None
            lines = content.split("\n")
            for line in lines[:10]:
                if line.startswith("# "):
                    title = line[2:].strip()
                    break

            chapters.append({
                "filename": chapter_file.name,
                "title": title,
                "content": content
            })
        except Exception as e:
            print(f"⚠️  Warning: Could not read {chapter_file}: {e}")

    return chapters


def get_image_map(repo_root: Path):
    """Create a map of available images for automatic insertion."""
    visuals_dir = repo_root / "series" / "hockey-shuttle" / "10-visuals"

    return {
        "episode_header": visuals_dir / "episode-headers" / "ep01-returning-to-center-ice.png",
        "sophia_portrait": visuals_dir / "characters" / "sophia-chen-athlete-v1.png",
        "ethan_portrait": visuals_dir / "characters" / "ethan-price-hockey-v1.png",
        "parking_lot_reunion": visuals_dir / "key-scenes" / "ep01-parking-lot-reunion.png",
        "winnipeg_winter": visuals_dir / "key-scenes" / "the-forks-winnipeg-winter.png",
        "empty_rink": visuals_dir / "atmospheric" / "empty-rink-golden-hour.png",
        "frost_window": visuals_dir / "atmospheric" / "frost-on-window.png",
        "snow_falling": visuals_dir / "atmospheric" / "snow-falling-night.png",
        "ice_texture": visuals_dir / "atmospheric" / "ice-texture-closeup.png",
        "truck_interior": visuals_dir / "atmospheric" / "truck-interior-dashboard.png",
        "coffee_cup": visuals_dir / "atmospheric" / "coffee-cup-winter.png",
        "shuttlecock": visuals_dir / "atmospheric" / "shuttlecock-in-flight.png",
        "hockey_puck": visuals_dir / "atmospheric" / "hockey-puck-closeup.png",
        "crossed_equipment": visuals_dir / "atmospheric" / "crossed-equipment.png",
        "two_paths": visuals_dir / "atmospheric" / "two-paths-snow.png",
    }


def get_chapter_images(chapter_num: int, image_map: dict):
    """Get images to insert for a specific chapter based on IMAGE-PLACEMENT-GUIDE.md logic."""
    images = []

    if chapter_num == 1:
        # Chapter 1: Episode header + childhood/ice images
        if image_map.get("episode_header") and image_map["episode_header"].exists():
            images.append(("opening", image_map["episode_header"], ""))
        if image_map.get("ice_texture") and image_map["ice_texture"].exists():
            images.append(("middle", image_map["ice_texture"], "Eight years ago..."))
        if image_map.get("empty_rink") and image_map["empty_rink"].exists():
            images.append(("end", image_map["empty_rink"], "Six years apart..."))

    elif chapter_num == 2:
        # Chapter 2: Winnipeg winter + parking lot reunion
        if image_map.get("winnipeg_winter") and image_map["winnipeg_winter"].exists():
            images.append(("opening", image_map["winnipeg_winter"], "Winnipeg, present day"))
        if image_map.get("snow_falling") and image_map["snow_falling"].exists():
            images.append(("before_reunion", image_map["snow_falling"], ""))
        if image_map.get("parking_lot_reunion") and image_map["parking_lot_reunion"].exists():
            images.append(("reunion", image_map["parking_lot_reunion"], "The moment everything changed"))
        if image_map.get("frost_window") and image_map["frost_window"].exists():
            images.append(("end", image_map["frost_window"], "One familiar face in a city of strangers"))

    elif chapter_num == 3:
        # Chapter 3: Badminton/truck/coffee
        if image_map.get("shuttlecock") and image_map["shuttlecock"].exists():
            images.append(("opening", image_map["shuttlecock"], ""))
        if image_map.get("truck_interior") and image_map["truck_interior"].exists():
            images.append(("middle", image_map["truck_interior"], "Later that evening..."))
        if image_map.get("coffee_cup") and image_map["coffee_cup"].exists():
            images.append(("end", image_map["coffee_cup"], "Warming up to something new"))

    elif chapter_num == 4:
        # Chapter 4: Hockey atmosphere
        if image_map.get("empty_rink") and image_map["empty_rink"].exists():
            images.append(("opening", image_map["empty_rink"], ""))
        if image_map.get("hockey_puck") and image_map["hockey_puck"].exists():
            images.append(("before_game", image_map["hockey_puck"], "Game night"))
        if image_map.get("crossed_equipment") and image_map["crossed_equipment"].exists():
            images.append(("end", image_map["crossed_equipment"], "Two worlds colliding"))

    return images


def create_chapter_header(chapter_num: int, chapter_title: str, episode_info: dict):
    """Create header text for chapter."""
    header_lines = [
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━",
        f"{episode_info['series']} - S{episode_info['season']}E{episode_info['episode']}",
        f"Chapter {chapter_num}",
    ]
    if chapter_title:
        header_lines.append(chapter_title)
    header_lines.append("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    return "\n".join(header_lines)


def create_chapter_footer(is_last_chapter=False):
    """Create footer text for chapter."""
    if is_last_chapter:
        return """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Thank you for reading!

Don't forget to vote ⭐ if you enjoyed this episode!

What did you think? Leave a comment! 💬

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""
    else:
        return """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

To be continued...

Don't forget to vote ⭐ and comment 💬

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""


def add_formatted_text(paragraph, text):
    """Add text to paragraph with markdown formatting (bold, italic)."""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic within non-bold text
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    # Italic text
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    # Regular text
                    if ipart:
                        paragraph.add_run(ipart)


def generate_chapter_docx(chapter: dict, chapter_num: int, episode_info: dict,
                          output_path: Path, image_map: dict, is_last: bool):
    """Generate DOCX file for a single chapter with formatting and images."""
    if not DOCX_AVAILABLE:
        print("⚠️  python-docx not installed. Install with: uv pip install python-docx")
        return False

    try:
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add chapter header
        header_text = create_chapter_header(chapter_num, chapter["title"], episode_info)
        header_para = doc.add_paragraph(header_text)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(11)

        doc.add_paragraph()  # Spacing

        # Get images for this chapter
        chapter_images = get_chapter_images(chapter_num, image_map)

        # Add opening image if available
        opening_images = [img for img in chapter_images if img[0] == "opening"]
        for _, img_path, caption in opening_images:
            try:
                doc.add_picture(str(img_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.runs[0].italic = True
                doc.add_paragraph()  # Spacing
            except Exception as e:
                print(f"⚠️  Could not add image {img_path}: {e}")

        # Parse chapter content
        lines = chapter["content"].split("\n")
        in_paragraph = False
        current_para = None
        skip_first_title = True
        line_count = 0

        for line in lines:
            line = line.rstrip()
            line_count += 1

            # Skip the main chapter title (already in header)
            if skip_first_title and line.startswith("# "):
                skip_first_title = False
                continue

            # Headings (convert to bold text)
            if line.startswith("## "):
                para = doc.add_paragraph()
                run = para.add_run(line[3:].strip())
                run.bold = True
                run.font.size = Pt(14)
                doc.add_paragraph()  # Spacing
                in_paragraph = False
                current_para = None
                continue
            elif line.startswith("### "):
                para = doc.add_paragraph()
                run = para.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(12)
                doc.add_paragraph()  # Spacing
                in_paragraph = False
                current_para = None
                continue

            # Scene breaks (horizontal rules)
            if line.strip() == "---":
                # Add middle images at scene breaks
                middle_images = [img for img in chapter_images if img[0] in ["middle", "before_reunion", "before_game"]]
                if middle_images and line_count > len(lines) / 3:  # Only after some content
                    for _, img_path, caption in middle_images[:1]:  # Add one at a time
                        try:
                            doc.add_paragraph()
                            doc.add_picture(str(img_path), width=Inches(6))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if caption:
                                caption_para = doc.add_paragraph(caption)
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.runs[0].italic = True
                            chapter_images = [img for img in chapter_images if img != (_, img_path, caption)]
                        except Exception as e:
                            print(f"⚠️  Could not add image {img_path}: {e}")

                # Scene break
                para = doc.add_paragraph("* * *")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacing
                in_paragraph = False
                current_para = None
                continue

            # Empty lines
            if not line.strip():
                in_paragraph = False
                current_para = None
                continue

            # Regular text with formatting
            if not in_paragraph or current_para is None:
                current_para = doc.add_paragraph()
                current_para.paragraph_format.line_spacing = 1.5
                in_paragraph = True
            else:
                current_para.add_run(" ")

            # Add text with bold/italic support
            add_formatted_text(current_para, line)

        # Add end images
        doc.add_paragraph()  # Spacing before end images
        end_images = [img for img in chapter_images if img[0] in ["end", "reunion"]]
        for _, img_path, caption in end_images:
            try:
                doc.add_picture(str(img_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.runs[0].italic = True
                doc.add_paragraph()  # Spacing
            except Exception as e:
                print(f"⚠️  Could not add image {img_path}: {e}")

        # Add footer
        footer_text = create_chapter_footer(is_last)
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        return True

    except Exception as e:
        print(f"❌ Error generating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main execution function."""
    args = parse_args()

    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx not installed")
        print("   Install with: uv pip install python-docx")
        sys.exit(1)

    repo_root = find_repo_root()
    episode_path = repo_root / args.episode_path

    if not episode_path.exists():
        print(f"❌ Error: Episode path not found: {episode_path}")
        sys.exit(1)

    episode_info = extract_episode_info(episode_path)

    print(f"\n🏒 Processing: {episode_info['episode_slug']}")
    print(f"   {episode_info['full_title']}\n")

    chapters = read_chapter_files(episode_path)

    if not chapters:
        sys.exit(1)

    print(f"📚 Found {len(chapters)} chapters\n")

    # Get image map
    image_map = get_image_map(repo_root)

    # Output directory
    output_dir = repo_root / args.output_dir / episode_info['episode_slug']
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"📝 Creating DOCX files for Wattpad...\n")

    for i, chapter in enumerate(chapters, 1):
        is_last = (i == len(chapters))

        # Generate chapter DOCX
        output_file = output_dir / f"chapter-{i:02d}.docx"

        if generate_chapter_docx(chapter, i, episode_info, output_file, image_map, is_last):
            chapter_name = chapter["title"] if chapter["title"] else f"Chapter {i}"
            print(f"✅ Saved: {output_file.name} - {chapter_name}")

    print(f"\n✨ Done! DOCX files ready for Wattpad!")
    print(f"📁 Location: {output_dir}\n")
    print("💡 To use:")
    print("   1. Open DOCX file in Microsoft Word")
    print("   2. Select All (Cmd+A / Ctrl+A)")
    print("   3. Copy (Cmd+C / Ctrl+C)")
    print("   4. Paste into Wattpad chapter editor")
    print("   5. Formatting (bold, italic) and images preserved!\n")


if __name__ == "__main__":
    main()
