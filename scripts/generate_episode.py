#!/usr/bin/env python3
"""
Episode Generator for The Boy Who Knew Me First Series

Generates PDF, HTML, EPUB, and DOCX formats from episode markdown files
with automatic inline image insertion and professional cover pages.

Usage:
    # For HTML, EPUB, DOCX (default)
    uv run scripts/generate_episode.py <episode-path>

    # For PDF generation (requires library path on macOS)
    export DYLD_LIBRARY_PATH=$(brew --prefix)/lib:$DYLD_LIBRARY_PATH
    uv run scripts/generate_episode.py <episode-path> --formats html,epub,docx,pdf

Examples:
    # Generate default formats (HTML, EPUB, DOCX)
    uv run scripts/generate_episode.py series/hockey-shuttle/season-01/episode-01-returning-to-center-ice

    # Generate all formats including PDF
    export DYLD_LIBRARY_PATH=$(brew --prefix)/lib:$DYLD_LIBRARY_PATH
    uv run scripts/generate_episode.py series/hockey-shuttle/season-01/episode-01-returning-to-center-ice --formats html,epub,docx,pdf

    # Generate only HTML and DOCX
    uv run scripts/generate_episode.py series/hockey-shuttle/season-01/episode-01-returning-to-center-ice --formats html,docx

Note:
    - PDF generation requires WeasyPrint with system libraries (pango, glib)
    - On macOS, set DYLD_LIBRARY_PATH before running for PDF support
    - Cover page uses empty-rink-golden-hour.png as background
    - Author is set to "Joxy" in all formats
"""

import argparse
import sys
from pathlib import Path
import markdown2
import re
from datetime import datetime

# Check for optional dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from weasyprint import HTML
    PDF_AVAILABLE = True
except (ImportError, OSError):
    PDF_AVAILABLE = False

try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False


def parse_args():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Generate PDF, HTML, EPUB, and DOCX from episode markdown files with images"
    )
    parser.add_argument(
        "episode_path",
        type=str,
        help="Path to episode directory"
    )
    parser.add_argument(
        "--formats",
        type=str,
        default="html,epub,docx",
        help="Comma-separated list of formats (html,epub,docx,pdf). Default: html,epub,docx"
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default="output",
        help="Output directory. Default: output"
    )
    parser.add_argument(
        "--with-images",
        action="store_true",
        default=True,
        help="Include images in output (default: True)"
    )
    return parser.parse_args()


def find_repo_root():
    """Find the repository root directory."""
    current = Path.cwd()
    while current != current.parent:
        if (current / ".git").exists() or (current / "series").exists():
            return current
        current = current.parent
    return Path.cwd()


def extract_episode_info(episode_path: Path):
    """Extract episode information from path."""
    episode_name = episode_path.name
    match = re.match(r"episode-(\d+)-(.+)", episode_name)

    if match:
        episode_num = int(match.group(1))
        episode_title = match.group(2).replace("-", " ").title()
    else:
        episode_num = "1"
        episode_title = episode_name

    return {
        "series": "The Boy Who Knew Me First",
        "season": "1",
        "episode": episode_num,
        "title": episode_title,
        "full_title": f"The Boy Who Knew Me First - Season 1, Episode {episode_num}: {episode_title}",
        "episode_slug": episode_name
    }


def read_chapter_files(episode_path: Path):
    """Read all chapter markdown files."""
    draft_dir = episode_path / "draft"

    if not draft_dir.exists():
        print(f"❌ Error: Draft directory not found at {draft_dir}")
        return []

    chapter_files = sorted(draft_dir.glob("chapter-*.md"))

    if not chapter_files:
        print(f"❌ Error: No chapter files found in {draft_dir}")
        return []

    chapters = []
    for chapter_file in chapter_files:
        try:
            content = chapter_file.read_text()

            # Extract chapter title
            title = None
            lines = content.split("\n")
            for line in lines[:10]:
                if line.startswith("# "):
                    title = line[2:].strip()
                    break

            chapters.append({
                "filename": chapter_file.name,
                "title": title,
                "content": content
            })
        except Exception as e:
            print(f"⚠️  Warning: Could not read {chapter_file}: {e}")

    return chapters


def get_image_map(repo_root: Path):
    """Create a map of available images for automatic insertion."""
    visuals_dir = repo_root / "series" / "hockey-shuttle" / "10-visuals"

    return {
        "episode_header": visuals_dir / "episode-headers" / "ep01-returning-to-center-ice.png",
        "sophia_portrait": visuals_dir / "characters" / "sophia-chen-athlete-v1.png",
        "ethan_portrait": visuals_dir / "characters" / "ethan-price-hockey-v1.png",
        "parking_lot_reunion": visuals_dir / "key-scenes" / "ep01-parking-lot-reunion.png",
        "winnipeg_winter": visuals_dir / "key-scenes" / "the-forks-winnipeg-winter.png",
        "empty_rink": visuals_dir / "atmospheric" / "empty-rink-golden-hour.png",
        "frost_window": visuals_dir / "atmospheric" / "frost-on-window.png",
        "snow_falling": visuals_dir / "atmospheric" / "snow-falling-night.png",
        "ice_texture": visuals_dir / "atmospheric" / "ice-texture-closeup.png",
        "truck_interior": visuals_dir / "atmospheric" / "truck-interior-dashboard.png",
        "coffee_cup": visuals_dir / "atmospheric" / "coffee-cup-winter.png",
        "shuttlecock": visuals_dir / "atmospheric" / "shuttlecock-in-flight.png",
        "hockey_puck": visuals_dir / "atmospheric" / "hockey-puck-closeup.png",
        "crossed_equipment": visuals_dir / "atmospheric" / "crossed-equipment.png",
        "two_paths": visuals_dir / "atmospheric" / "two-paths-snow.png",
    }


def combine_chapters(chapters: list, episode_info: dict, include_images: bool, repo_root: Path):
    """Combine all chapters into single markdown content with optional images."""
    combined = []

    # Title page
    combined.append(f"# {episode_info['full_title']}\n")
    combined.append("---\n")

    image_map = get_image_map(repo_root) if include_images else {}

    for i, chapter in enumerate(chapters, 1):
        # Chapter heading
        if chapter["title"]:
            combined.append(f"## Chapter {i}: {chapter['title']}\n")
        else:
            combined.append(f"## Chapter {i}\n")

        # Add episode header image at start of chapter 1
        if i == 1 and include_images and image_map.get("episode_header"):
            img_path = image_map["episode_header"]
            if img_path.exists():
                combined.append(f"\n![Episode Header]({img_path})\n")

        # Add chapter content
        content = chapter["content"]
        lines = content.split("\n")

        # Skip first line if it's a title
        if lines and lines[0].startswith("# "):
            lines = lines[1:]

        combined.append("\n".join(lines).strip())
        combined.append("")

        # Add chapter-specific images based on content/position
        if include_images:
            if i == 2 and image_map.get("parking_lot_reunion"):
                # Add reunion scene in chapter 2
                img_path = image_map["parking_lot_reunion"]
                if img_path.exists():
                    combined.append(f"\n![The Reunion]({img_path})\n")

        combined.append("---\n")

    return "\n".join(combined)


def generate_html(markdown_content: str, episode_info: dict, output_path: Path, cover_image_path: Path = None):
    """Generate HTML file from markdown with images and cover page."""
    html_content = markdown2.markdown(
        markdown_content,
        extras=['fenced-code-blocks', 'tables', 'break-on-newline']
    )

    # Create cover page HTML if cover image exists
    cover_html = ""
    if cover_image_path and cover_image_path.exists():
        cover_html = f"""
    <div class="cover-page">
        <div class="cover-content">
            <h1 class="cover-title">The Boy Who Knew Me First</h1>
            <p class="cover-season">Season {episode_info['season']}</p>
            <p class="cover-episode">Episode {episode_info['episode']}</p>
            <p class="cover-author">by Joxy</p>
        </div>
    </div>
    <div style="page-break-after: always;"></div>
    """

    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{episode_info['full_title']}</title>
    <style>
        body {{
            font-family: Georgia, 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 0;
            color: #333;
        }}
        .cover-page {{
            width: 100%;
            height: 100vh;
            background-image: url('{cover_image_path}');
            background-size: cover;
            background-position: center center;
            background-repeat: no-repeat;
            display: flex;
            align-items: center;
            justify-content: center;
            position: relative;
            margin: 0;
            padding: 0;
            overflow: hidden;
        }}
        .cover-page::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: rgba(0, 0, 0, 0.4);
        }}
        .cover-content {{
            position: relative;
            z-index: 1;
            text-align: center;
            color: white;
            text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.9);
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
        }}
        .cover-title {{
            font-size: 4em;
            margin: 0 0 0.5em 0;
            font-weight: 700;
            letter-spacing: 3px;
            text-transform: uppercase;
            color: #FFFFFF;
        }}
        .cover-season {{
            font-size: 1.5em;
            margin: 0.5em 0;
            font-weight: 400;
            color: #FFFFFF;
        }}
        .cover-episode {{
            font-size: 1.5em;
            margin: 0.5em 0;
            font-weight: 400;
            color: #FFFFFF;
        }}
        .cover-author {{
            font-size: 1.2em;
            margin: 2em 0 0 0;
            font-style: italic;
            font-weight: 300;
            color: #FFFFFF;
        }}
        .content {{
            padding: 20px;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 2em auto;
            border-radius: 4px;
        }}
        hr {{
            border: none;
            border-top: 1px solid #bdc3c7;
            margin: 40px 0;
        }}
        @media print {{
            body {{ max-width: 100%; }}
            .cover-page {{ page-break-after: always; }}
        }}
    </style>
</head>
<body>
{cover_html}
<div class="content">
{html_content}
</div>
</body>
</html>"""

    output_path.write_text(html_doc, encoding='utf-8')
    return True


def generate_docx(markdown_content: str, episode_info: dict, output_path: Path, image_map: dict, cover_image_path: Path = None):
    """Generate DOCX file with inline images and cover page."""
    if not DOCX_AVAILABLE:
        print("⚠️  python-docx not installed. Install with: uv pip install python-docx")
        return False

    try:
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

        # Cover page with background image
        if cover_image_path and cover_image_path.exists():
            try:
                # Import PIL for image manipulation
                from PIL import Image

                # Open and crop image to fill page (8.5 x 11 inches)
                img = Image.open(cover_image_path)
                page_width_px = int(8.5 * 300)  # 8.5 inches at 300 DPI
                page_height_px = int(11 * 300)  # 11 inches at 300 DPI
                page_ratio = page_width_px / page_height_px
                img_ratio = img.width / img.height

                # Crop to fill entire page (center crop)
                if img_ratio > page_ratio:
                    # Image is wider - crop width
                    new_width = int(img.height * page_ratio)
                    left = (img.width - new_width) // 2
                    img_cropped = img.crop((left, 0, left + new_width, img.height))
                else:
                    # Image is taller - crop height
                    new_height = int(img.width / page_ratio)
                    top = (img.height - new_height) // 2
                    img_cropped = img.crop((0, top, img.width, top + new_height))

                # Save cropped image temporarily
                import tempfile
                temp_cover = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                img_cropped.save(temp_cover.name, 'PNG')

                # Add full-page background image (fills entire page)
                doc.add_picture(temp_cover.name, width=Inches(8.5), height=Inches(11))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Clean up temp file
                import os
                os.unlink(temp_cover.name)

                # Add spacing to position text in center
                for _ in range(8):
                    doc.add_paragraph()

                # Title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run('The Boy Who Knew Me First')
                title_run.font.size = Pt(48)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(255, 255, 255)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Season
                season_para = doc.add_paragraph()
                season_run = season_para.add_run(f'Season {episode_info["season"]}')
                season_run.font.size = Pt(24)
                season_run.font.color.rgb = RGBColor(255, 255, 255)
                season_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Episode
                episode_para = doc.add_paragraph()
                episode_run = episode_para.add_run(f'Episode {episode_info["episode"]}')
                episode_run.font.size = Pt(24)
                episode_run.font.color.rgb = RGBColor(255, 255, 255)
                episode_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Author (with more spacing)
                for _ in range(4):
                    doc.add_paragraph()
                author_para = doc.add_paragraph()
                author_run = author_para.add_run('by Joxy')
                author_run.font.size = Pt(18)
                author_run.font.italic = True
                author_run.font.color.rgb = RGBColor(255, 255, 255)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            except Exception as e:
                print(f"⚠️  Could not add cover image: {e}")
                import traceback
                traceback.print_exc()

        # Page break after cover
        doc.add_page_break()

        # Reset margins for content
        sections = doc.sections
        for section in sections[1:]:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Parse markdown and add to document
        lines = markdown_content.split("\n")
        in_paragraph = False
        current_para = None

        for line in lines:
            line = line.rstrip()

            # Check for image references
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = Path(img_match.group(2))
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Add caption
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption.runs[0].italic = True
                        doc.add_paragraph()  # Spacing
                    except Exception as e:
                        print(f"⚠️  Could not add image {img_path}: {e}")
                continue

            # Headings
            if line.startswith("# "):
                doc.add_heading(line[2:], 1)
                in_paragraph = False
            elif line.startswith("## "):
                doc.add_heading(line[3:], 2)
                in_paragraph = False
            elif line.startswith("### "):
                doc.add_heading(line[4:], 3)
                in_paragraph = False
            # Horizontal rules
            elif line.strip() == "---":
                doc.add_page_break()
                in_paragraph = False
            # Empty lines
            elif not line.strip():
                in_paragraph = False
                current_para = None
            # Regular text
            else:
                if not in_paragraph or current_para is None:
                    current_para = doc.add_paragraph()
                    in_paragraph = True
                else:
                    current_para.add_run(" ")

                # Simple markdown parsing for bold and italic
                text = line
                # Handle bold
                text = re.sub(r'\*\*([^\*]+)\*\*', lambda m: m.group(1), text)
                # Handle italic
                text = re.sub(r'\*([^\*]+)\*', lambda m: m.group(1), text)

                run = current_para.add_run(text)

                # Apply styles based on original markdown
                if "**" in line:
                    run.bold = True
                if line.startswith("*") and not line.startswith("**"):
                    run.italic = True

        doc.save(output_path)
        return True

    except Exception as e:
        print(f"❌ Error generating DOCX: {e}")
        return False


def generate_epub(markdown_content: str, episode_info: dict, output_path: Path, cover_image_path: Path = None):
    """Generate EPUB file from markdown content with cover page."""
    if not EPUB_AVAILABLE:
        print("⚠️  ebooklib not installed. Install with: uv pip install ebooklib")
        return False

    try:
        # Create EPUB book
        book = epub.EpubBook()

        # Set metadata
        book.set_identifier(f"hockey-shuttle-s{episode_info['season']}-e{episode_info['episode']}")
        book.set_title(episode_info['full_title'])
        book.set_language('en')
        book.add_author('Joxy')

        # Add cover image if available
        if cover_image_path and cover_image_path.exists():
            try:
                from PIL import Image, ImageDraw, ImageFont
                import tempfile
                import os

                # Crop image to standard e-book ratio (1600x2560, 5:8 ratio)
                img = Image.open(cover_image_path)
                ebook_width = 1600
                ebook_height = 2560
                ebook_ratio = ebook_width / ebook_height
                img_ratio = img.width / img.height

                # Center crop to fill
                if img_ratio > ebook_ratio:
                    # Image is wider - crop width
                    new_width = int(img.height * ebook_ratio)
                    left = (img.width - new_width) // 2
                    img_cropped = img.crop((left, 0, left + new_width, img.height))
                else:
                    # Image is taller - crop height
                    new_height = int(img.width / ebook_ratio)
                    top = (img.height - new_height) // 2
                    img_cropped = img.crop((0, top, img.width, top + new_height))

                # Resize to standard e-book dimensions
                img_with_bg = img_cropped.resize((ebook_width, ebook_height), Image.Resampling.LANCZOS)

                # Draw text overlay directly on the image
                draw = ImageDraw.Draw(img_with_bg)

                # Try to use Orkney font by Hanken Design Co
                import os
                home = os.path.expanduser("~")
                try:
                    # Orkney fonts for EPUB cover
                    author_font = ImageFont.truetype(f"{home}/Library/Fonts/orkney-medium.otf", 140)
                    title_font = ImageFont.truetype(f"{home}/Library/Fonts/orkney-bold.otf", 240)
                    subtitle_font = ImageFont.truetype(f"{home}/Library/Fonts/orkney-regular.otf", 110)
                except:
                    try:
                        # Fallback to Arial
                        author_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 140)
                        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 240)
                        subtitle_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 110)
                    except:
                        try:
                            # Fallback to Helvetica
                            author_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 140)
                            title_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 240)
                            subtitle_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 110)
                        except:
                            author_font = ImageFont.load_default()
                            title_font = ImageFont.load_default()
                            subtitle_font = ImageFont.load_default()

                # Calculate center positions
                center_x = ebook_width // 2

                # Text content
                author_text = "Joxy"
                title_line1 = "HOCKEY"
                title_line2 = "SHUTTLE"
                season_text = f"Season {episode_info['season']}"
                episode_text = f"Episode {episode_info['episode']}"

                # Orange/gold color for title
                title_color = (255, 180, 50)  # Orange/gold

                # Helper function to draw text with shadow
                def draw_text_with_shadow(text, font, y_pos, color=(255, 255, 255), is_title=False):
                    # Get text bounding box
                    bbox = draw.textbbox((0, 0), text, font=font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    x_pos = center_x - text_width // 2

                    # Draw shadow (black, slightly offset)
                    shadow_offset = 10 if is_title else 6
                    draw.text((x_pos + shadow_offset, y_pos + shadow_offset), text,
                             font=font, fill=(0, 0, 0, 230))

                    # Draw main text
                    draw.text((x_pos, y_pos), text, font=font, fill=color)

                    return text_height

                # Author at the very top
                author_y = 120
                draw_text_with_shadow(author_text, author_font, author_y, color=(255, 255, 255))

                # Title, Season, Episode near the bottom
                bottom_start = ebook_height - 800  # Start 800px from bottom

                # Title line 1: HOCKEY
                title1_height = draw_text_with_shadow(title_line1, title_font, bottom_start,
                                                     color=title_color, is_title=True)

                # Title line 2: SHUTTLE
                title2_y = bottom_start + title1_height + 20
                title2_height = draw_text_with_shadow(title_line2, title_font, title2_y,
                                                     color=title_color, is_title=True)

                # Season
                season_y = title2_y + title2_height + 80
                season_height = draw_text_with_shadow(season_text, subtitle_font, season_y,
                                                     color=(255, 255, 255))

                # Episode
                episode_y = season_y + season_height + 50
                draw_text_with_shadow(episode_text, subtitle_font, episode_y,
                                     color=(255, 255, 255))

                # Save to temp file
                temp_cover = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                img_with_bg.save(temp_cover.name, 'PNG')

                # Add to EPUB
                with open(temp_cover.name, 'rb') as cover_file:
                    book.set_cover('cover.png', cover_file.read())

                # Clean up
                os.unlink(temp_cover.name)

            except Exception as e:
                print(f"⚠️  Could not process cover image: {e}")
                import traceback
                traceback.print_exc()
                # Fallback to original image
                with open(cover_image_path, 'rb') as cover_file:
                    book.set_cover('cover.png', cover_file.read())

            # Create cover page HTML
            cover_html = f'''
            <html>
            <head>
                <style>
                    body {{
                        margin: 0;
                        padding: 0;
                        text-align: center;
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
                    }}
                    .cover {{
                        position: relative;
                        width: 100%;
                        height: 100vh;
                        overflow: hidden;
                    }}
                    .cover img {{
                        width: 100%;
                        height: 100%;
                        object-fit: cover;
                        object-position: center center;
                    }}
                    .cover-text {{
                        position: absolute;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        color: #FFFFFF;
                        text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.9);
                        width: 100%;
                    }}
                    h1 {{
                        font-size: 3em;
                        margin: 0.5em 0;
                        font-weight: 700;
                        letter-spacing: 3px;
                        text-transform: uppercase;
                        color: #FFFFFF;
                    }}
                    p {{
                        font-size: 1.5em;
                        margin: 0.3em 0;
                        font-weight: 400;
                        color: #FFFFFF;
                    }}
                    .author {{
                        font-style: italic;
                        margin-top: 2em;
                        font-weight: 300;
                        color: #FFFFFF;
                    }}
                </style>
            </head>
            <body>
                <div class="cover">
                    <img src="cover.png" alt="Cover"/>
                    <div class="cover-text">
                        <h1>The Boy Who Knew Me First</h1>
                        <p>Season {episode_info['season']}</p>
                        <p>Episode {episode_info['episode']}</p>
                        <p class="author">by Joxy</p>
                    </div>
                </div>
            </body>
            </html>
            '''

            cover_page = epub.EpubHtml(title='Cover', file_name='cover.xhtml', lang='en')
            cover_page.content = cover_html
            book.add_item(cover_page)

        # Split into chapters
        chapters_content = markdown_content.split("---")
        epub_chapters = []

        for i, chapter_content in enumerate(chapters_content):
            if not chapter_content.strip():
                continue

            html_content = markdown2.markdown(chapter_content, extras=['fenced-code-blocks', 'tables'])

            # Create EPUB chapter
            c = epub.EpubHtml(
                title=f"Chapter {i}",
                file_name=f"chap_{i:02d}.xhtml",
                lang='en'
            )
            c.content = f"<html><body>{html_content}</body></html>"

            book.add_item(c)
            epub_chapters.append(c)

        # Add default NCX and Nav files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Define CSS style
        style = '''
        body { font-family: Georgia, serif; line-height: 1.6; }
        h1 { color: #2c3e50; }
        p { text-align: justify; margin: 1em 0; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)

        # Create spine (include cover if it exists)
        if cover_image_path and cover_image_path.exists():
            book.spine = ['cover', 'nav'] + epub_chapters
        else:
            book.spine = ['nav'] + epub_chapters

        # Write EPUB file
        epub.write_epub(output_path, book, {})
        return True

    except Exception as e:
        print(f"❌ Error generating EPUB: {e}")
        return False


def generate_pdf(html_path: Path, output_path: Path):
    """Generate PDF file from HTML."""
    if not PDF_AVAILABLE:
        print("⚠️  WeasyPrint not available. Install with: uv pip install weasyprint")
        return False

    try:
        HTML(filename=str(html_path)).write_pdf(output_path)
        return True
    except Exception as e:
        print(f"❌ Error generating PDF: {e}")
        return False


def main():
    """Main execution function."""
    args = parse_args()

    repo_root = find_repo_root()
    episode_path = repo_root / args.episode_path

    if not episode_path.exists():
        print(f"❌ Error: Episode path not found: {episode_path}")
        sys.exit(1)

    # Extract episode info
    episode_info = extract_episode_info(episode_path)

    print(f"\n📚 Processing: {episode_info['episode_slug']}")
    print(f"   {episode_info['full_title']}\n")

    # Read chapters
    chapters = read_chapter_files(episode_path)

    if not chapters:
        sys.exit(1)

    print(f"✓ Found {len(chapters)} chapters\n")

    # Get image map
    image_map = get_image_map(repo_root) if args.with_images else {}

    # Get cover image path
    cover_image = repo_root / "series" / "hockey-shuttle" / "10-visuals" / "atmospheric" / "empty-rink-golden-hour.png"

    # Combine chapters with images
    markdown_content = combine_chapters(chapters, episode_info, args.with_images, repo_root)

    # Set up output directory
    output_dir = repo_root / args.output_dir / episode_info['episode_slug']
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate requested formats
    formats = [f.strip().lower() for f in args.formats.split(',')]
    generated_files = []

    for fmt in formats:
        output_file = output_dir / f"{episode_info['episode_slug']}.{fmt}"

        if fmt == 'html':
            if generate_html(markdown_content, episode_info, output_file, cover_image):
                print(f"✅ Generated HTML: {output_file}")
                generated_files.append(output_file)

        elif fmt == 'docx' or fmt == 'doc':
            if DOCX_AVAILABLE:
                if generate_docx(markdown_content, episode_info, output_file.with_suffix('.docx'), image_map, cover_image):
                    print(f"✅ Generated DOCX: {output_file.with_suffix('.docx')}")
                    generated_files.append(output_file.with_suffix('.docx'))
            else:
                print(f"⚠️  Skipping DOCX: python-docx not installed")
                print(f"   Install with: uv pip install python-docx")

        elif fmt == 'epub':
            if EPUB_AVAILABLE:
                if generate_epub(markdown_content, episode_info, output_file, cover_image):
                    print(f"✅ Generated EPUB: {output_file}")
                    generated_files.append(output_file)
            else:
                print(f"⚠️  Skipping EPUB: ebooklib not installed")
                print(f"   Install with: uv pip install ebooklib")

        elif fmt == 'pdf':
            # Generate HTML first if not already done
            html_file = output_dir / f"{episode_info['episode_slug']}.html"
            if not html_file.exists():
                generate_html(markdown_content, episode_info, html_file, cover_image)

            if PDF_AVAILABLE:
                if generate_pdf(html_file, output_file):
                    print(f"✅ Generated PDF: {output_file}")
                    generated_files.append(output_file)
            else:
                print(f"⚠️  Skipping PDF: WeasyPrint not available")

        else:
            print(f"⚠️  Unknown format: {fmt}")

    print(f"\n✨ Done! Generated {len(generated_files)} files")
    print(f"📁 Output directory: {output_dir}\n")


if __name__ == "__main__":
    main()
