#!/usr/bin/env python3
"""
Create Aesthetics and Playlist DOCX for Wattpad

Generates a beautiful intro chapter with character aesthetic grids and curated playlist.

Usage:
    uv run scripts/create_aesthetics_playlist.py
"""

from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def find_repo_root():
    """Find the repository root directory."""
    current = Path.cwd()
    while current != current.parent:
        if (current / ".git").exists() or (current / "series").exists():
            return current
        current = current.parent
    return Path.cwd()


def create_aesthetics_playlist_docx():
    """Create the Aesthetics and Playlist DOCX file."""
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx not installed")
        print("   Install with: uv pip install python-docx")
        return False

    repo_root = find_repo_root()

    # Locate aesthetic grids
    sophia_grid = repo_root / "series/hockey-shuttle/10-visuals/aesthetics/sophia-aesthetic-grid.png"
    ethan_grid = repo_root / "series/hockey-shuttle/10-visuals/aesthetics/ethan-aesthetic-grid.png"

    if not sophia_grid.exists() or not ethan_grid.exists():
        print(f"❌ Error: Aesthetic grid images not found")
        return False

    try:
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title = doc.add_heading('The Boy Who Knew Me First', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph('Season 1 • Aesthetics & Playlist')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].italic = True

        doc.add_paragraph()  # Spacing

        # Introduction
        intro = doc.add_paragraph()
        intro_run = intro.add_run('Welcome to The Boy Who Knew Me First! Before we dive into the story, meet our main characters through their aesthetics and the soundtrack that defines their journey.')
        intro_run.italic = True
        intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        # ═══════════════════════════════════════
        # SOPHIA CHEN
        # ═══════════════════════════════════════

        divider1 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sophia_heading = doc.add_heading('SOPHIA CHEN', 1)
        sophia_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sophia_tagline = doc.add_paragraph()
        tagline_run = sophia_tagline.add_run('The New Girl | Badminton Player | Physics Nerd')
        tagline_run.italic = True
        sophia_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

        divider2 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Sophia bio
        sophia_bio = doc.add_paragraph()
        sophia_bio.add_run('Age: ').bold = True
        sophia_bio.add_run('17\n')
        sophia_bio.add_run('Sport: ').bold = True
        sophia_bio.add_run('Badminton (provincial level)\n')
        sophia_bio.add_run('Passion: ').bold = True
        sophia_bio.add_run('Physics, astronomy, understanding how things work\n')
        sophia_bio.add_run('Vibe: ').bold = True
        sophia_bio.add_run('Teal and winter aesthetics, cozy study sessions, quiet determination\n\n')

        sophia_bio.add_run('Just moved to Winnipeg for senior year. New city, new school, zero friends—except one. The kid who taught her to skate when they were nine. The one she hasn\'t seen in six years.\n\n')

        sophia_bio.add_run('Athletic but not obsessed. Smart but not pretentious. She traded hockey for badminton years ago, but some connections never really fade.')
        sophia_bio.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        # Sophia aesthetic grid
        doc.add_picture(str(sophia_grid), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption1 = doc.add_paragraph('Sophia\'s Aesthetic Board')
        caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption1.runs[0].italic = True
        caption1.runs[0].font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        # ═══════════════════════════════════════
        # ETHAN PRICE
        # ═══════════════════════════════════════

        divider3 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ethan_heading = doc.add_heading('ETHAN PRICE', 1)
        ethan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ethan_tagline = doc.add_paragraph()
        tagline_run2 = ethan_tagline.add_run('Hockey Captain | Mr. Popular | The One Who Got Away')
        tagline_run2.italic = True
        ethan_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

        divider4 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Ethan bio
        ethan_bio = doc.add_paragraph()
        ethan_bio.add_run('Age: ').bold = True
        ethan_bio.add_run('17\n')
        ethan_bio.add_run('Sport: ').bold = True
        ethan_bio.add_run('Hockey (AAA team captain)\n')
        ethan_bio.add_run('Secret: ').bold = True
        ethan_bio.add_run('Still loves physics (just doesn\'t tell anyone)\n')
        ethan_bio.add_run('Vibe: ').bold = True
        ethan_bio.add_run('Navy and ice rink aesthetics, leadership, hidden depths\n\n')

        ethan_bio.add_run('The popular hockey captain everyone knows. Six years of training transformed the kid who taught Sophia to skate into someone she barely recognizes—except when he smiles, and suddenly he\'s nine years old again.\n\n')

        ethan_bio.add_run('He has a girlfriend. He has a future in hockey. He has everything planned out.\n\n')

        ethan_bio.add_run('And then Sophia Chen walks back into his life.')
        ethan_bio.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        # Ethan aesthetic grid
        doc.add_picture(str(ethan_grid), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption2 = doc.add_paragraph('Ethan\'s Aesthetic Board')
        caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption2.runs[0].italic = True
        caption2.runs[0].font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        # ═══════════════════════════════════════
        # PLAYLIST
        # ═══════════════════════════════════════

        divider5 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        playlist_heading = doc.add_heading('THE SOUNDTRACK', 1)
        playlist_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        playlist_subtitle = doc.add_paragraph()
        subtitle_run = playlist_subtitle.add_run('Songs That Tell Their Story')
        subtitle_run.italic = True
        playlist_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        divider6 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Playlist intro
        playlist_intro = doc.add_paragraph()
        intro_text = playlist_intro.add_run('Listen along as you read! Here\'s the playlist that captures Sophia and Ethan\'s journey from childhood friends to... something more.')
        intro_text.italic = True
        playlist_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Playlist tracks
        playlist = [
            ('1. "Teenage Dream" – Olivia Rodrigo', 'The nostalgia of what once was'),
            ('2. "Skating" – Vance Joy', 'For the memories on ice'),
            ('3. "The Night We Met" – Lord Huron', 'Six years apart, one moment to reconnect'),
            ('4. "Slow Dancing in a Burning Room" – John Mayer', 'When things get complicated'),
            ('5. "Drivers License" – Olivia Rodrigo', 'Watching him with someone else'),
            ('6. "Sofia" – Clairo', 'Because her name is Sophia (and it fits perfectly)'),
            ('7. "Heather" – Conan Gray', 'She\'s everything I wish I could be'),
            ('8. "Falling" – Harry Styles', 'When the walls come down'),
            ('9. "Electric Love" – BØRNS', 'That spark that won\'t go away'),
            ('10. "Can\'t Help Falling in Love" – Kina Grannis', 'Slow burn at its finest'),
            ('11. "Moral of the Story" – Ashe', 'Life lessons and growing up'),
            ('12. "Dark Paradise" – Lana Del Rey', 'Winter vibes, Winnipeg nights'),
            ('13. "Golden" – Harry Styles', 'He\'s everything golden'),
            ('14. "Ivy" – Taylor Swift', 'Tangled up in forbidden feelings'),
            ('15. "The 1" – Taylor Swift', 'What if timing was different?'),
        ]

        for track, description in playlist:
            track_para = doc.add_paragraph()
            track_run = track_para.add_run(track)
            track_run.bold = True
            track_para.add_run('\n     ')
            desc_run = track_para.add_run(description)
            desc_run.italic = True
            track_para.paragraph_format.line_spacing = 1.3

        doc.add_paragraph()
        doc.add_paragraph()

        # ═══════════════════════════════════════
        # FOOTER / CALL TO ACTION
        # ═══════════════════════════════════════

        divider7 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider7.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Story tagline
        tagline_para = doc.add_paragraph()
        tagline1 = tagline_para.add_run('🏒 Childhood Best Friends → Six Years Apart → One Last Chance')
        tagline1.bold = True
        tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        tagline2_para = doc.add_paragraph()
        tagline2_para.add_run('✨ He has a girlfriend\n')
        tagline2_para.add_run('✨ She\'s not trying to be ').italic = True
        that_girl = tagline2_para.add_run('that')
        that_girl.italic = True
        that_girl.bold = True
        tagline2_para.add_run(' girl\n').italic = True
        tagline2_para.add_run('✨ Neither of them can stop thinking about the other')
        tagline2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline2_para.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        finale = doc.add_paragraph()
        finale_text = finale.add_run('Ten episodes. One season. A slow burn so intense it\'ll melt the ice. 🔥')
        finale_text.bold = True
        finale.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        divider8 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Call to action
        cta = doc.add_paragraph()
        cta.add_run('Ready to start reading?\n\n').bold = True
        cta.add_run('⭐ Vote if you love the vibe!\n')
        cta.add_run('💬 Comment which song you\'re most excited to match with a scene!\n')
        cta.add_run('🔔 Follow for episode updates!')
        cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cta.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        divider9 = doc.add_paragraph('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        divider9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        next_chapter = doc.add_paragraph()
        next_text = next_chapter.add_run('Next: Episode 1, Chapter 1 - "Then"')
        next_text.italic = True
        next_chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save
        output_dir = repo_root / "output" / "wattpad-docx"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / "00-aesthetics-and-playlist.docx"

        doc.save(output_file)

        print(f"\n✅ Created: {output_file}")
        print(f"\n📊 Contents:")
        print(f"   • Character aesthetic grids (Sophia & Ethan)")
        print(f"   • Character bios and descriptions")
        print(f"   • 15-song curated playlist")
        print(f"   • Story taglines and hooks")
        print(f"\n💡 This should be your FIRST chapter on Wattpad!")
        print(f"   Title it: 'Aesthetics & Playlist' or '00 - Meet the Characters'\n")

        return True

    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main execution."""
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx not installed")
        print("   Install with: uv pip install python-docx")
        sys.exit(1)

    create_aesthetics_playlist_docx()


if __name__ == "__main__":
    main()
